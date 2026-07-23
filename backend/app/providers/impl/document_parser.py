import logging
import re

import aiofiles
from pypdf import PdfReader
from starlette.concurrency import run_in_threadpool

from app.providers.base import DocumentParser
from app.providers.types import ParsedDocument, ParsedPage

logger = logging.getLogger(__name__)

SUPPORTED_MIMES = {
    "application/pdf",
    "text/plain",
    "text/markdown",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = re.sub(r"\r\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


class LlamaDocumentParser(DocumentParser):
    """Extract text from PDF, TXT, MD, and DOCX files."""

    async def parse(self, file_path: str, mime_type: str) -> ParsedDocument:
        if mime_type == "application/pdf":
            return await self._parse_pdf(file_path)
        if mime_type in ("text/plain", "text/markdown"):
            return await self._parse_text(file_path)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return await self._parse_docx(file_path)
        raise ValueError(f"Unsupported file type: {mime_type}")

    async def _parse_text(self, file_path: str) -> ParsedDocument:
        async with aiofiles.open(file_path, encoding="utf-8", errors="replace") as f:
            text = _clean_text(await f.read())
        return ParsedDocument(pages=[ParsedPage(page_number=1, text=text)], full_text=text)

    async def _parse_pdf(self, file_path: str) -> ParsedDocument:
        def _read() -> ParsedDocument:
            reader = PdfReader(file_path)
            pages: list[ParsedPage] = []
            parts: list[str] = []
            for i, page in enumerate(reader.pages, start=1):
                raw = page.extract_text() or ""
                cleaned = _clean_text(raw)
                if cleaned:
                    pages.append(ParsedPage(page_number=i, text=cleaned))
                    parts.append(cleaned)
            full = "\n\n".join(parts)
            return ParsedDocument(pages=pages, full_text=full)

        return await run_in_threadpool(_read)

    async def _parse_docx(self, file_path: str) -> ParsedDocument:
        def _read() -> ParsedDocument:
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            paragraphs = [_clean_text(p.text) for p in doc.paragraphs if p.text.strip()]
            full = "\n\n".join(paragraphs)
            return ParsedDocument(pages=[ParsedPage(page_number=1, text=full)], full_text=full)

        return await run_in_threadpool(_read)
